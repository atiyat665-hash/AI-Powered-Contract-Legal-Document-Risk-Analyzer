"""
Document Text Extraction Service
Supports PDF, DOCX, TXT, and OCR for scanned documents
"""
import os
import re
import pdfplumber
import pytesseract
from docx import Document
from PIL import Image
from io import BytesIO


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF using pdfplumber.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text from all pages
    """
    try:
        text_content = []
        
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
        
        full_text = "\n\n".join(text_content)
        return full_text
        
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX using python-docx.
    Extracts text from paragraphs and tables.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text from document
    """
    try:
        doc = Document(file_path)
        text_content = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        full_text = "\n".join(text_content)
        return full_text
        
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return ""


def extract_text_from_txt(file_path: str) -> str:
    """
    Extract text from plain text file.
    Tries UTF-8 first, falls back to latin-1 if needed.
    
    Args:
        file_path: Path to the TXT file
        
    Returns:
        Extracted text from file
    """
    try:
        # Try UTF-8 first
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Fallback to latin-1
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
                
    except Exception as e:
        print(f"Error reading text file: {e}")
        return ""


def extract_text_from_scanned_pdf(file_path: str) -> str:
    """
    Extract text from scanned PDF using OCR (pytesseract).
    Uses pdfplumber to convert pages to images, then applies OCR.
    This is a fallback method for PDFs with little or no extractable text.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        OCR-extracted text from all pages
        
    Note:
        Requires Tesseract OCR to be installed on the system.
        Windows users: Install from https://github.com/UB-Mannheim/tesseract/wiki
    """
    try:
        text_content = []
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                try:
                    # Convert page to image
                    page_image = page.to_image(resolution=300)
                    
                    # Get PIL Image
                    pil_image = page_image.original
                    
                    # Apply OCR
                    ocr_text = pytesseract.image_to_string(pil_image)
                    
                    if ocr_text.strip():
                        text_content.append(f"--- Page {page_num + 1} ---\n{ocr_text}")
                        
                except Exception as page_error:
                    print(f"Error processing page {page_num + 1}: {page_error}")
                    continue
        
        full_text = "\n\n".join(text_content)
        return full_text
        
    except Exception as e:
        print(f"Error performing OCR on PDF: {e}")
        return ""


def clean_text(text: str) -> str:
    """
    Clean extracted text by removing excessive whitespace,
    fixing line breaks, and stripping weird characters.
    
    Args:
        text: Raw extracted text
        
    Returns:
        Cleaned text
    """
    if not text:
        return ""
    
    # Remove null bytes and other control characters (except newline, tab, carriage return)
    text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]', '', text)
    
    # Normalize line breaks (convert \r\n and \r to \n)
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    
    # Remove excessive blank lines (more than 2 consecutive newlines)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Remove trailing/leading whitespace from each line
    lines = [line.rstrip() for line in text.split('\n')]
    text = '\n'.join(lines)
    
    # Remove excessive spaces (more than 2 consecutive spaces)
    text = re.sub(r' {3,}', '  ', text)
    
    # Remove leading/trailing whitespace from entire text
    text = text.strip()
    
    return text


def extract_text(file_path: str, file_type: str) -> tuple[str, bool]:
    """
    Main extraction function that routes to the appropriate extractor
    based on file type, applies text cleaning, and returns final text.
    
    Args:
        file_path: Path to the uploaded file
        file_type: File extension (pdf, docx, txt)
        
    Returns:
        Tuple of (extracted_text, used_ocr)
        - extracted_text: Cleaned text content
        - used_ocr: Boolean indicating if OCR was used
        
    Raises:
        ValueError: If file_type is not supported
    """
    file_type = file_type.lower().strip('.')
    used_ocr = False
    
    try:
        # Route to appropriate extractor
        if file_type == 'pdf':
            # Try normal PDF extraction first
            raw_text = extract_text_from_pdf(file_path)
            
            # If extraction returned very little text, try OCR
            if len(raw_text.strip()) < 50:
                print(f"PDF appears to be scanned (only {len(raw_text)} chars). Attempting OCR...")
                ocr_text = extract_text_from_scanned_pdf(file_path)
                
                # Use OCR text if it's better than normal extraction
                if len(ocr_text) > len(raw_text):
                    raw_text = ocr_text
                    used_ocr = True
                    print(f"OCR successful: extracted {len(ocr_text)} characters")
        
        elif file_type == 'docx':
            raw_text = extract_text_from_docx(file_path)
        
        elif file_type == 'txt':
            raw_text = extract_text_from_txt(file_path)
        
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Clean the extracted text
        cleaned_text = clean_text(raw_text)
        
        return cleaned_text, used_ocr
        
    except Exception as e:
        print(f"Error during text extraction: {e}")
        raise


def get_text_stats(text: str) -> dict:
    """
    Get statistics about extracted text.
    
    Args:
        text: Extracted text
        
    Returns:
        Dictionary with text statistics
    """
    if not text:
        return {
            'characters': 0,
            'words': 0,
            'lines': 0,
            'paragraphs': 0
        }
    
    # Count characters (excluding whitespace)
    char_count = len(text.replace(' ', '').replace('\n', '').replace('\t', ''))
    
    # Count words
    word_count = len(text.split())
    
    # Count lines (non-empty)
    lines = [line for line in text.split('\n') if line.strip()]
    line_count = len(lines)
    
    # Count paragraphs (separated by blank lines)
    paragraphs = [p for p in text.split('\n\n') if p.strip()]
    paragraph_count = len(paragraphs)
    
    return {
        'characters': char_count,
        'words': word_count,
        'lines': line_count,
        'paragraphs': paragraph_count
    }
