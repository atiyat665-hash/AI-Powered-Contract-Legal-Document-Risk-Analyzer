"""
Report generation service.
Generates professional PDF and DOCX reports summarizing contract analysis,
risk assessment, and executive summary for a document.
"""
import os
import json
from datetime import datetime

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
)

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPORTS_DIR = "reports"
os.makedirs(REPORTS_DIR, exist_ok=True)

RISK_COLORS = {
    "High": colors.HexColor("#e74c3c"),
    "Medium": colors.HexColor("#f39c12"),
    "Low": colors.HexColor("#f1c40f"),
}


def _safe(value, default="N/A"):
    """Return value if truthy, otherwise a default placeholder."""
    return value if value else default


def generate_pdf_report(document, analysis, risks_list, summary):
    """Generate a professional PDF report and return its file path."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"report_{document.id}_{timestamp}.pdf"
    filepath = os.path.join(REPORTS_DIR, filename)

    doc = SimpleDocTemplate(
        filepath, pagesize=A4,
        topMargin=2 * cm, bottomMargin=2 * cm,
        leftMargin=2 * cm, rightMargin=2 * cm
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('TitleStyle', parent=styles['Title'], fontSize=20, spaceAfter=6)
    heading_style = ParagraphStyle('HeadingStyle', parent=styles['Heading2'], spaceBefore=14, spaceAfter=8)
    normal = styles['Normal']

    elements = []

    # Header
    elements.append(Paragraph("AI Contract Risk Analysis Report", title_style))
    elements.append(Paragraph(f"Filename: {_safe(document.filename)}", normal))
    elements.append(Paragraph(f"Generated: {datetime.now().strftime('%d %b %Y, %I:%M %p')}", normal))
    elements.append(Spacer(1, 12))

    # Section 1: Document Info
    elements.append(Paragraph("1. Document Information", heading_style))
    doc_info = [
        ["Filename", _safe(document.filename)],
        ["File Type", _safe(document.file_type)],
        ["Upload Date", str(_safe(document.upload_date))],
    ]
    t = Table(doc_info, colWidths=[5 * cm, 10 * cm])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor("#f0f0f0")),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    elements.append(t)

    # Section 2: Contract Analysis
    elements.append(Paragraph("2. Contract Analysis", heading_style))
    if analysis:
        analysis_rows = [
            ["Contract Type", _safe(analysis.contract_type)],
            ["Parties", _safe(analysis.parties)],
            ["Effective Date", _safe(analysis.effective_date)],
            ["Expiry Date", _safe(analysis.expiry_date)],
            ["Payment Terms", _safe(analysis.payment_terms)],
            ["Renewal Clause", _safe(analysis.renewal_clause)],
            ["Termination Clause", _safe(analysis.termination_clause)],
            ["Confidentiality", _safe(analysis.confidentiality)],
            ["Jurisdiction", _safe(analysis.jurisdiction)],
        ]
        t2 = Table(analysis_rows, colWidths=[5 * cm, 10 * cm])
        t2.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor("#f0f0f0")),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        elements.append(t2)
    else:
        elements.append(Paragraph("No analysis data available.", normal))

    # Section 3: Risk Assessment
    elements.append(Paragraph("3. Risk Assessment", heading_style))
    weights = {"High": 3, "Medium": 2, "Low": 1}
    if risks_list:
        total = sum(weights.get(r.risk_level, 1) for r in risks_list)
        score = min(int((total / (len(risks_list) * 3)) * 100), 100)
    else:
        score = 0
    elements.append(Paragraph(f"<b>Overall Risk Score: {score}/100</b>", normal))
    elements.append(Spacer(1, 8))

    if risks_list:
        risk_table_data = [["Risk Title", "Level", "Confidence", "Explanation", "Recommendation"]]
        for r in risks_list:
            risk_table_data.append([
                Paragraph(_safe(r.risk_title), normal),
                r.risk_level,
                f"{int((r.confidence_score or 0) * 100)}%",
                Paragraph(_safe(r.explanation), normal),
                Paragraph(_safe(r.recommendation), normal),
            ])
        t3 = Table(risk_table_data, colWidths=[3.5 * cm, 1.8 * cm, 2 * cm, 4 * cm, 4 * cm])
        style_cmds = [
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#34495e")),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]
        for i, r in enumerate(risks_list, start=1):
            color = RISK_COLORS.get(r.risk_level, colors.grey)
            style_cmds.append(('BACKGROUND', (1, i), (1, i), color))
        t3.setStyle(TableStyle(style_cmds))
        elements.append(t3)
    else:
        elements.append(Paragraph("No significant risks detected.", normal))

    # Section 4: Executive Summary
    elements.append(Paragraph("4. Executive Summary", heading_style))
    if summary:
        elements.append(Paragraph(_safe(summary.executive_summary), normal))
        elements.append(Spacer(1, 6))
        elements.append(Paragraph(f"<b>Payment Summary:</b> {_safe(summary.payment_summary)}", normal))
        elements.append(Paragraph(f"<b>Termination Summary:</b> {_safe(summary.termination_summary)}", normal))
        elements.append(Paragraph(f"<b>Risk Summary:</b> {_safe(summary.risk_summary)}", normal))
    else:
        elements.append(Paragraph("No summary available.", normal))

    # Section 5: Recommended Actions
    elements.append(Paragraph("5. Recommended Actions", heading_style))
    if summary and summary.recommended_actions:
        try:
            actions = json.loads(summary.recommended_actions)
        except Exception:
            actions = [summary.recommended_actions]
        for i, action in enumerate(actions, 1):
            elements.append(Paragraph(f"{i}. {action}", normal))
    else:
        elements.append(Paragraph("No recommendations available.", normal))

    elements.append(Spacer(1, 20))
    elements.append(Paragraph("Generated by AI Contract Risk Analyzer", styles['Italic']))

    doc.build(elements)
    return filepath


def generate_docx_report(document, analysis, risks_list, summary):
    """Generate a professional DOCX report and return its file path."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"report_{document.id}_{timestamp}.docx"
    filepath = os.path.join(REPORTS_DIR, filename)

    d = DocxDocument()

    d.add_heading("AI Contract Risk Analysis Report", level=0)
    d.add_paragraph(f"Filename: {_safe(document.filename)}")
    d.add_paragraph(f"Generated: {datetime.now().strftime('%d %b %Y, %I:%M %p')}")

    d.add_heading("1. Document Information", level=1)
    d.add_paragraph(f"Filename: {_safe(document.filename)}")
    d.add_paragraph(f"File Type: {_safe(document.file_type)}")
    d.add_paragraph(f"Upload Date: {_safe(document.upload_date)}")

    d.add_heading("2. Contract Analysis", level=1)
    if analysis:
        fields = [
            ("Contract Type", analysis.contract_type),
            ("Parties", analysis.parties),
            ("Effective Date", analysis.effective_date),
            ("Expiry Date", analysis.expiry_date),
            ("Payment Terms", analysis.payment_terms),
            ("Renewal Clause", analysis.renewal_clause),
            ("Termination Clause", analysis.termination_clause),
            ("Confidentiality", analysis.confidentiality),
            ("Jurisdiction", analysis.jurisdiction),
        ]
        for label, value in fields:
            p = d.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(_safe(value))
    else:
        d.add_paragraph("No analysis data available.")

    d.add_heading("3. Risk Assessment", level=1)
    weights = {"High": 3, "Medium": 2, "Low": 1}
    if risks_list:
        total = sum(weights.get(r.risk_level, 1) for r in risks_list)
        score = min(int((total / (len(risks_list) * 3)) * 100), 100)
    else:
        score = 0
    p = d.add_paragraph()
    p.add_run(f"Overall Risk Score: {score}/100").bold = True

    if risks_list:
        table = d.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = (
            "Risk Title", "Level", "Confidence", "Explanation", "Recommendation"
        )
        for r in risks_list:
            row = table.add_row().cells
            row[0].text = _safe(r.risk_title)
            row[1].text = r.risk_level
            row[2].text = f"{int((r.confidence_score or 0) * 100)}%"
            row[3].text = _safe(r.explanation)
            row[4].text = _safe(r.recommendation)
    else:
        d.add_paragraph("No significant risks detected.")

    d.add_heading("4. Executive Summary", level=1)
    if summary:
        d.add_paragraph(_safe(summary.executive_summary))
        p = d.add_paragraph()
        p.add_run("Payment Summary: ").bold = True
        p.add_run(_safe(summary.payment_summary))
        p = d.add_paragraph()
        p.add_run("Termination Summary: ").bold = True
        p.add_run(_safe(summary.termination_summary))
        p = d.add_paragraph()
        p.add_run("Risk Summary: ").bold = True
        p.add_run(_safe(summary.risk_summary))
    else:
        d.add_paragraph("No summary available.")

    d.add_heading("5. Recommended Actions", level=1)
    if summary and summary.recommended_actions:
        try:
            actions = json.loads(summary.recommended_actions)
        except Exception:
            actions = [summary.recommended_actions]
        for action in actions:
            d.add_paragraph(action, style='List Number')
    else:
        d.add_paragraph("No recommendations available.")

    footer = d.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Generated by AI Contract Risk Analyzer")
    run.italic = True

    d.save(filepath)
    return filepath
    """
Report generation service.
Generates professional PDF and DOCX reports summarizing contract analysis,
risk assessment, and executive summary for a document.
"""
import os
import json
from datetime import datetime

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
)

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPORTS_DIR = "reports"
os.makedirs(REPORTS_DIR, exist_ok=True)

RISK_COLORS = {
    "High": colors.HexColor("#e74c3c"),
    "Medium": colors.HexColor("#f39c12"),
    "Low": colors.HexColor("#f1c40f"),
}


def _safe(value, default="N/A"):
    """Return value if truthy, otherwise a default placeholder."""
    return value if value else default


def generate_pdf_report(document, analysis, risks_list, summary):
    """Generate a professional PDF report and return its file path."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"report_{document.id}_{timestamp}.pdf"
    filepath = os.path.join(REPORTS_DIR, filename)

    doc = SimpleDocTemplate(
        filepath, pagesize=A4,
        topMargin=2 * cm, bottomMargin=2 * cm,
        leftMargin=2 * cm, rightMargin=2 * cm
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('TitleStyle', parent=styles['Title'], fontSize=20, spaceAfter=6)
    heading_style = ParagraphStyle('HeadingStyle', parent=styles['Heading2'], spaceBefore=14, spaceAfter=8)
    normal = styles['Normal']

    elements = []

    # Header
    elements.append(Paragraph("AI Contract Risk Analysis Report", title_style))
    elements.append(Paragraph(f"Filename: {_safe(document.filename)}", normal))
    elements.append(Paragraph(f"Generated: {datetime.now().strftime('%d %b %Y, %I:%M %p')}", normal))
    elements.append(Spacer(1, 12))

    # Section 1: Document Info
    elements.append(Paragraph("1. Document Information", heading_style))
    doc_info = [
        ["Filename", _safe(document.filename)],
        ["File Type", _safe(document.file_type)],
        ["Upload Date", str(_safe(document.upload_date))],
    ]
    t = Table(doc_info, colWidths=[5 * cm, 10 * cm])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor("#f0f0f0")),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    elements.append(t)

    # Section 2: Contract Analysis
    elements.append(Paragraph("2. Contract Analysis", heading_style))
    if analysis:
        analysis_rows = [
            ["Contract Type", _safe(analysis.contract_type)],
            ["Parties", _safe(analysis.parties)],
            ["Effective Date", _safe(analysis.effective_date)],
            ["Expiry Date", _safe(analysis.expiry_date)],
            ["Payment Terms", _safe(analysis.payment_terms)],
            ["Renewal Clause", _safe(analysis.renewal_clause)],
            ["Termination Clause", _safe(analysis.termination_clause)],
            ["Confidentiality", _safe(analysis.confidentiality)],
            ["Jurisdiction", _safe(analysis.jurisdiction)],
        ]
        t2 = Table(analysis_rows, colWidths=[5 * cm, 10 * cm])
        t2.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor("#f0f0f0")),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        elements.append(t2)
    else:
        elements.append(Paragraph("No analysis data available.", normal))

    # Section 3: Risk Assessment
    elements.append(Paragraph("3. Risk Assessment", heading_style))
    weights = {"High": 3, "Medium": 2, "Low": 1}
    if risks_list:
        total = sum(weights.get(r.risk_level, 1) for r in risks_list)
        score = min(int((total / (len(risks_list) * 3)) * 100), 100)
    else:
        score = 0
    elements.append(Paragraph(f"<b>Overall Risk Score: {score}/100</b>", normal))
    elements.append(Spacer(1, 8))

    if risks_list:
        risk_table_data = [["Risk Title", "Level", "Confidence", "Explanation", "Recommendation"]]
        for r in risks_list:
            risk_table_data.append([
                Paragraph(_safe(r.risk_title), normal),
                r.risk_level,
                f"{int((r.confidence_score or 0) * 100)}%",
                Paragraph(_safe(r.explanation), normal),
                Paragraph(_safe(r.recommendation), normal),
            ])
        t3 = Table(risk_table_data, colWidths=[3.5 * cm, 1.8 * cm, 2 * cm, 4 * cm, 4 * cm])
        style_cmds = [
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#34495e")),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]
        for i, r in enumerate(risks_list, start=1):
            color = RISK_COLORS.get(r.risk_level, colors.grey)
            style_cmds.append(('BACKGROUND', (1, i), (1, i), color))
        t3.setStyle(TableStyle(style_cmds))
        elements.append(t3)
    else:
        elements.append(Paragraph("No significant risks detected.", normal))

    # Section 4: Executive Summary
    elements.append(Paragraph("4. Executive Summary", heading_style))
    if summary:
        elements.append(Paragraph(_safe(summary.executive_summary), normal))
        elements.append(Spacer(1, 6))
        elements.append(Paragraph(f"<b>Payment Summary:</b> {_safe(summary.payment_summary)}", normal))
        elements.append(Paragraph(f"<b>Termination Summary:</b> {_safe(summary.termination_summary)}", normal))
        elements.append(Paragraph(f"<b>Risk Summary:</b> {_safe(summary.risk_summary)}", normal))
    else:
        elements.append(Paragraph("No summary available.", normal))

    # Section 5: Recommended Actions
    elements.append(Paragraph("5. Recommended Actions", heading_style))
    if summary and summary.recommended_actions:
        try:
            actions = json.loads(summary.recommended_actions)
        except Exception:
            actions = [summary.recommended_actions]
        for i, action in enumerate(actions, 1):
            elements.append(Paragraph(f"{i}. {action}", normal))
    else:
        elements.append(Paragraph("No recommendations available.", normal))

    elements.append(Spacer(1, 20))
    elements.append(Paragraph("Generated by AI Contract Risk Analyzer", styles['Italic']))

    doc.build(elements)
    return filepath


def generate_docx_report(document, analysis, risks_list, summary):
    """Generate a professional DOCX report and return its file path."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"report_{document.id}_{timestamp}.docx"
    filepath = os.path.join(REPORTS_DIR, filename)

    d = DocxDocument()

    d.add_heading("AI Contract Risk Analysis Report", level=0)
    d.add_paragraph(f"Filename: {_safe(document.filename)}")
    d.add_paragraph(f"Generated: {datetime.now().strftime('%d %b %Y, %I:%M %p')}")

    d.add_heading("1. Document Information", level=1)
    d.add_paragraph(f"Filename: {_safe(document.filename)}")
    d.add_paragraph(f"File Type: {_safe(document.file_type)}")
    d.add_paragraph(f"Upload Date: {_safe(document.upload_date)}")

    d.add_heading("2. Contract Analysis", level=1)
    if analysis:
        fields = [
            ("Contract Type", analysis.contract_type),
            ("Parties", analysis.parties),
            ("Effective Date", analysis.effective_date),
            ("Expiry Date", analysis.expiry_date),
            ("Payment Terms", analysis.payment_terms),
            ("Renewal Clause", analysis.renewal_clause),
            ("Termination Clause", analysis.termination_clause),
            ("Confidentiality", analysis.confidentiality),
            ("Jurisdiction", analysis.jurisdiction),
        ]
        for label, value in fields:
            p = d.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(_safe(value))
    else:
        d.add_paragraph("No analysis data available.")

    d.add_heading("3. Risk Assessment", level=1)
    weights = {"High": 3, "Medium": 2, "Low": 1}
    if risks_list:
        total = sum(weights.get(r.risk_level, 1) for r in risks_list)
        score = min(int((total / (len(risks_list) * 3)) * 100), 100)
    else:
        score = 0
    p = d.add_paragraph()
    p.add_run(f"Overall Risk Score: {score}/100").bold = True

    if risks_list:
        table = d.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = (
            "Risk Title", "Level", "Confidence", "Explanation", "Recommendation"
        )
        for r in risks_list:
            row = table.add_row().cells
            row[0].text = _safe(r.risk_title)
            row[1].text = r.risk_level
            row[2].text = f"{int((r.confidence_score or 0) * 100)}%"
            row[3].text = _safe(r.explanation)
            row[4].text = _safe(r.recommendation)
    else:
        d.add_paragraph("No significant risks detected.")

    d.add_heading("4. Executive Summary", level=1)
    if summary:
        d.add_paragraph(_safe(summary.executive_summary))
        p = d.add_paragraph()
        p.add_run("Payment Summary: ").bold = True
        p.add_run(_safe(summary.payment_summary))
        p = d.add_paragraph()
        p.add_run("Termination Summary: ").bold = True
        p.add_run(_safe(summary.termination_summary))
        p = d.add_paragraph()
        p.add_run("Risk Summary: ").bold = True
        p.add_run(_safe(summary.risk_summary))
    else:
        d.add_paragraph("No summary available.")

    d.add_heading("5. Recommended Actions", level=1)
    if summary and summary.recommended_actions:
        try:
            actions = json.loads(summary.recommended_actions)
        except Exception:
            actions = [summary.recommended_actions]
        for action in actions:
            d.add_paragraph(action, style='List Number')
    else:
        d.add_paragraph("No recommendations available.")

    footer = d.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Generated by AI Contract Risk Analyzer")
    run.italic = True

    d.save(filepath)
    return filepath